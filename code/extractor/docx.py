import docx


class DOCXReader:
    @staticmethod
    def extract(filepath: str) -> str:
        """Returns DOCX pages content."""
        doc = docx.Document(filepath)
        text = ""

        # Read available tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text

        # Read available paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text

        return text
